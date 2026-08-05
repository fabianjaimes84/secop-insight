#!/usr/bin/env python3
"""
Script para examinar la plantilla Word y ver qué campos/placeholders tiene.
"""

from docx import Document
import re

PLANTILLA_PATH = "plantillas/Formato 1 - Carta de presentación de la oferta.docx"

def examinar_plantilla():
    try:
        doc = Document(PLANTILLA_PATH)

        print(f"[*] Examinando plantilla: {PLANTILLA_PATH}\n")
        print(f"[*] Total de párrafos: {len(doc.paragraphs)}\n")

        # Buscar placeholders (texto entre {{ }})
        placeholders = set()

        print("=" * 80)
        print("CONTENIDO DE LA PLANTILLA")
        print("=" * 80 + "\n")

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[P{i}] {para.text[:100]}...")

                # Buscar placeholders
                matches = re.findall(r'\{\{(.+?)\}\}', para.text)
                if matches:
                    for match in matches:
                        placeholders.add(match.strip())
                        print(f"     → Placeholder encontrado: {{{{{match}}}}}")

        print("\n" + "=" * 80)
        print("RESUMEN DE PLACEHOLDERS")
        print("=" * 80 + "\n")

        if placeholders:
            print(f"Total de placeholders únicos: {len(placeholders)}\n")
            for placeholder in sorted(placeholders):
                print(f"  - {{{{{placeholder}}}}}")
        else:
            print("No se encontraron placeholders en formato {{...}}")
            print("\nPrimer párrafo (para referencia):")
            if doc.paragraphs:
                print(f"  {doc.paragraphs[0].text}")

        # Examinar tablas
        if doc.tables:
            print(f"\n[*] Documento contiene {len(doc.tables)} tabla(s)")
            for i, table in enumerate(doc.tables):
                print(f"\n  Tabla {i+1}: {len(table.rows)} filas, {len(table.columns)} columnas")
                for j, row in enumerate(table.rows[:3]):  # Primeras 3 filas
                    cells_text = [cell.text[:30] for cell in row.cells]
                    print(f"    Fila {j}: {cells_text}")

    except Exception as e:
        print(f"\n[ERROR] {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    examinar_plantilla()
