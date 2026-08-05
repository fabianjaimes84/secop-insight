"""
Servicio para generar documentos a partir de plantillas.
Carga plantillas Word, las rellena con datos del proponente, y genera PDF.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path
import re
from typing import Optional
import subprocess
import tempfile
import os

PLANTILLAS_DIR = Path(__file__).parent.parent.parent / "plantillas"


def reemplazar_texto_en_documento(doc: Document, reemplazos: dict) -> None:
    """
    Reemplaza texto en párrafos y tablas del documento.
    Usa XML manipulation para garantizar que funcione correctamente.
    Maneja también tablas anidadas.

    Args:
        doc: Documento Word
        reemplazos: Dict con {texto_original: texto_nuevo}
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    # Función auxiliar para reemplazar en un párrafo
    def reemplazar_en_parrafo(para, reemplazos):
        for original, nuevo in reemplazos.items():
            if original in para.text:
                # Obtener el elemento XML del párrafo
                p = para._element

                # Obtener todo el texto del párrafo
                texto_completo = para.text

                if original in texto_completo:
                    # Reemplazar en el texto completo
                    texto_nuevo_completo = texto_completo.replace(original, nuevo)

                    # Eliminar todos los runs existentes
                    for run in para.runs:
                        r = run._element
                        r.getparent().remove(r)

                    # Crear un nuevo run con el texto reemplazado
                    new_run = para.add_run(texto_nuevo_completo)

                    print(f"    Reemplazo: '{original}' = '{nuevo[:30]}...'")

    # Reemplazar en párrafos principales
    for para in doc.paragraphs:
        reemplazar_en_parrafo(para, reemplazos)

    # Reemplazar en tablas (incluyendo anidadas)
    def procesar_tabla(table):
        for row in table.rows:
            for cell in row.cells:
                # Reemplazar en párrafos de la celda
                for para in cell.paragraphs:
                    reemplazar_en_parrafo(para, reemplazos)

                # Procesar tablas anidadas dentro de la celda
                for tabla_anidada in cell.tables:
                    procesar_tabla(tabla_anidada)

    # Procesar todas las tablas del documento
    for table in doc.tables:
        procesar_tabla(table)


def generar_carta_presentacion(
    nombre_proponente: str,
    numero_proceso: str,
    objeto_proceso: str,
    nombre_representante: str,
    cedula_representante: str,
    direccion: str,
    correo: str,
    telefono: str,
    ciudad: str,
    salida_path: Path,
    tipo_proponente: str = "union_temporal",
    pertenece_grupo: bool = False,
    tipo_grupo: str = "",
    cotiza_bolsa: bool = False,
    accionistas: list = None,
    entidad: str = "",
    direccion_ejecucion: str = "",
    lotes: str = "",
    matricula_profesional: str = "",
) -> bool:
    """
    Genera la Carta de Presentación de Oferta con datos específicos.

    Args:
        nombre_proponente: Nombre del proponente
        numero_proceso: Número de proceso SECOP
        objeto_proceso: Objeto del proceso
        nombre_representante: Nombre del representante legal
        cedula_representante: Cédula del representante
        direccion: Dirección
        correo: Correo electrónico
        telefono: Teléfono
        ciudad: Ciudad
        salida_path: Ruta donde guardar el documento

    Returns:
        True si se generó exitosamente, False si hubo error
    """
    try:
        plantilla_path = PLANTILLAS_DIR / "Formato 1 - Carta de presentación de la oferta.docx"

        if not plantilla_path.exists():
            print(f"[ERROR] Plantilla no encontrada: {plantilla_path}")
            return False

        print(f"\n[*] Generando Carta de Presentación")
        print(f"    Proponente: {nombre_proponente}")
        print(f"    Proceso: {numero_proceso}")
        print(f"    Representante: {nombre_representante}")

        # Cargar documento
        doc = Document(plantilla_path)
        print(f"[OK] Plantilla cargada")

        # Limpiar número de proceso (remover paréntesis)
        numero_proceso_limpio = numero_proceso.split(" (")[0] if " (" in numero_proceso else numero_proceso

        # Marcar tipo de proponente con X
        marca_natural = "X" if tipo_proponente == "natural" else ""
        marca_union = "X" if tipo_proponente == "union_temporal" else ""
        marca_consorcio = "X" if tipo_proponente == "consorcio" else ""

        # Marcar grupo empresarial y cotiza en bolsa
        marca_grupo = "X" if pertenece_grupo else ""
        marca_bolsa = "X" if cotiza_bolsa else ""

        # Obtener datos de accionistas (máximo 2)
        accionista_1_nombre = ""
        accionista_1_porcentaje = ""
        accionista_1_cedula = ""
        accionista_2_nombre = ""
        accionista_2_porcentaje = ""
        accionista_2_cedula = ""

        if accionistas:
            if len(accionistas) > 0:
                accionista_1_nombre = accionistas[0].get("nombre", "")
                accionista_1_porcentaje = accionistas[0].get("porcentaje", "")
                accionista_1_cedula = accionistas[0].get("cedula", "")
            if len(accionistas) > 1:
                accionista_2_nombre = accionistas[1].get("nombre", "")
                accionista_2_porcentaje = accionistas[1].get("porcentaje", "")
                accionista_2_cedula = accionistas[1].get("cedula", "")

        # Definir reemplazos usando los placeholders del documento
        reemplazos = {
            "«no_proceso»": numero_proceso_limpio,
            "«nom_proponente»": nombre_proponente,
            "«nom_pers_nat__o_rp_propo_1»": nombre_representante,
            "«repre_cedula»": cedula_representante,
            "«repre_direcc_corr»": direccion,
            "«repre_email»": correo,
            "«repre_tele»": telefono,
            "«repre_ciudad»": ciudad,
            "«repre_mat_profe»": matricula_profesional,
            "«nombre_entidad»": entidad,
            "«dir_entidad»": direccion_ejecucion,
            "«objeto»": objeto_proceso,
            "«lote»": lotes,
            "«propo_per_natu»": marca_natural,
            "«propo_union»": marca_union,
            "«propo_conso»": marca_consorcio,
            "«grupo_m»": marca_grupo,
            "«cot_bolsa»": marca_bolsa,
            "«nom_o_raz_soc_1»": accionista_1_nombre,
            "«nom_o_raz_soc_2»": accionista_2_nombre,
            "«porcen_partici_1»": accionista_1_porcentaje,
            "«cedula_accion_1»": accionista_1_cedula,
            "«nom_accion_1»": accionista_1_nombre,
            "«porcen_partici_2»": accionista_2_porcentaje,
            "«cedula_accion_2»": accionista_2_cedula,
            "«nom_accion_2»": accionista_2_nombre,
            "«contac_nombre»": nombre_representante,
            "«contac_dir_ciudad»": direccion,
            "«contac_tele»": telefono,
            "«contac_email»": correo,
        }

        print(f"[*] Reemplazando {len(reemplazos)} valores...")

        # Reemplazar en documento
        reemplazar_texto_en_documento(doc, reemplazos)
        print(f"[OK] Valores reemplazados")

        # Guardar documento
        salida_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(salida_path)

        print(f"[OK] Documento guardado: {salida_path}\n")
        return True

    except Exception as e:
        print(f"[ERROR] {e}")
        import traceback
        traceback.print_exc()
        return False


def convertir_word_a_pdf(word_path: Path, pdf_path: Path) -> bool:
    """
    Convierte un documento Word a PDF usando LibreOffice.

    Args:
        word_path: Ruta del archivo Word
        pdf_path: Ruta donde guardar el PDF

    Returns:
        True si se convirtió exitosamente
    """
    try:
        # Crear carpeta de salida si no existe
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        # Comando para convertir con LibreOffice
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(pdf_path.parent),
            str(word_path)
        ]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

        if result.returncode == 0:
            print(f"[OK] PDF generado: {pdf_path}")
            return True
        else:
            print(f"[ERROR] LibreOffice error: {result.stderr}")
            return False

    except FileNotFoundError:
        print("[ERROR] LibreOffice no está instalado")
        return False
    except Exception as e:
        print(f"[ERROR] {e}")
        return False


if __name__ == "__main__":
    # Prueba
    salida = Path("salida_test.docx")
    generar_carta_presentacion(
        nombre_proponente="NUEVA EMPRESA TEST",
        numero_proceso="SI-CMA-004-2026",
        objeto_proceso="Prueba de contratación",
        nombre_representante="Juan Pérez García",
        cedula_representante="123.456.789",
        direccion="Calle 100 No. 50-20",
        correo="juan@empresa.com",
        telefono="3101234567",
        ciudad="Bogotá",
        salida_path=salida,
    )
