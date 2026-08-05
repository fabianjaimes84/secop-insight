#!/usr/bin/env python3
"""
Examina el documento en detalle para encontrar los textos a reemplazar.
"""

from docx import Document
from pathlib import Path

PLANTILLA_PATH = "plantillas/Formato 1 - Carta de presentación de la oferta.docx"

def examinar():
    doc = Document(PLANTILLA_PATH)

    print("\n=== PÁRRAFOS CON DATOS IMPORTANTES ===\n")

    # Buscar párrafos que contienen los datos que queremos reemplazar
    palabras_clave = ["VIALTEK", "GUISA", "RAFAEL", "Bucaramanga", "contratacion"]

    for i, para in enumerate(doc.paragraphs):
        texto = para.text

        # Buscar párrafos relevantes
        if any(palabra in texto for palabra in palabras_clave):
            print(f"[P{i}]")
            print(f"  Texto: {texto}")
            print(f"  Largo: {len(texto)}")
            print(f"  Runs: {len(para.runs)}")

            # Mostrar cada run por separado
            for j, run in enumerate(para.runs):
                print(f"    Run[{j}]: '{run.text}' (len={len(run.text)})")
            print()

    print("\n=== TABLAS ===\n")
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTabla {table_idx}: {len(table.rows)} filas x {len(table.columns)} columnas")
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                texto = cell.text
                print(f"  [{row_idx}][{col_idx}]: {texto[:80]}")

                # Si contiene palabras clave, mostrar detalles
                if any(palabra in texto for palabra in palabras_clave):
                    print(f"    >>> MATCH: {texto}")
                    for para_idx, para in enumerate(cell.paragraphs):
                        print(f"      Para[{para_idx}]: {para.text}")

if __name__ == "__main__":
    examinar()
